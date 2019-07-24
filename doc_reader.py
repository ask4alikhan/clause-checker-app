Tesseract: 
This fetches the sections =>(Section.*?\\n\\n)
This fetches before + all sections w/positive lookahead => .*?(?=(Section.*?\\n\\n))
This fetches the after sections bit => \\n\\n\[.* \\n\\n\[.*


Notes from Docx bits: 
import docx
doc = docx.Document('kobe-sett-2001-04-27-shtml-file_html.docx')
len(doc.paragraphs)
print(doc.paragraphs[0].text)
doc = docx
for idx,p in enumerate(doc.paragraphs):
    print(idx, p.text)
----